#!/usr/bin/env python3
"""Convert a practical Markdown report into a conservative print-ready DOCX."""

import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TOTAL_DXA = 9360
C = {
    "navy": "17365D", "blue": "2E74B5", "dark": "1F4D78", "ink": "202124",
    "muted": "5F6368", "light": "E8EEF5", "gray": "F2F4F7", "red_bg": "FCE8E6",
    "red": "9B1C1C", "gold_bg": "FFF4CE", "gold": "7A5A00", "border": "C9D2DC",
}
INLINE = re.compile(r'(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`|\*([^*]+)\*)')


def font(run, size=11, bold=None, italic=None, color="202124", name="Calibri"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic


def hyperlink(p, label, url):
    rid = p.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), "Calibri"); rf.set(qn("w:hAnsi"), "Calibri")
    col = OxmlElement("w:color"); col.set(qn("w:val"), C["blue"])
    ul = OxmlElement("w:u"); ul.set(qn("w:val"), "single")
    rpr.extend([rf, col, ul]); run.append(rpr)
    text = OxmlElement("w:t"); text.text = label; run.append(text); link.append(run); p._p.append(link)


def inline(p, text, size=11, color=None):
    color = color or C["ink"]
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos: font(p.add_run(text[pos:m.start()]), size, color=color)
        if m.group(2): hyperlink(p, m.group(2), m.group(3))
        elif m.group(4): font(p.add_run(m.group(4)), size, bold=True, color=color)
        elif m.group(5): font(p.add_run(m.group(5)), size - .5, color=C["dark"], name="Consolas")
        else: font(p.add_run(m.group(6)), size, italic=True, color=color)
        pos = m.end()
    if pos < len(text): font(p.add_run(text[pos:]), size, color=color)


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); shd = pr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); pr.append(shd)
    shd.set(qn("w:fill"), fill)


def no_split(row):
    pr = row._tr.get_or_add_trPr(); el = OxmlElement("w:cantSplit"); pr.append(el)


def repeat_header(row):
    pr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); pr.append(el)


def cell_margins(cell):
    pr = cell._tc.get_or_add_tcPr(); mar = pr.first_child_found_in("w:tcMar")
    if mar is None: mar = OxmlElement("w:tcMar"); pr.append(mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), str(value)); el.set(qn("w:type"), "dxa"); mar.append(el)


def geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    for tag, value, typ in (("tblW", sum(widths), "dxa"), ("tblInd", 120, "dxa")):
        el = pr.find(qn(f"w:{tag}"))
        if el is None: el = OxmlElement(f"w:{tag}"); pr.append(el)
        el.set(qn("w:w"), str(value)); el.set(qn("w:type"), typ)
    layout = pr.find(qn("w:tblLayout"))
    if layout is None: layout = OxmlElement("w:tblLayout"); pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            pr = cell._tc.get_or_add_tcPr(); tcw = pr.find(qn("w:tcW"))
            if tcw is None: tcw = OxmlElement("w:tcW"); pr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i])); tcw.set(qn("w:type"), "dxa")
            cell_margins(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def widths(headers, rows):
    n = len(headers)
    if n == 1: return [TOTAL_DXA]
    scores = []
    numeric = ("valor", "saldo", "preço", "r$/m²", "dono", "ordem", "peso", "cenário", "área", "prioridade")
    for j, h in enumerate(headers):
        length = max([len(h)] + [len(r[j]) if j < len(r) else 0 for r in rows])
        score = min(max(length, 10), 65)
        if any(x in h.lower() for x in numeric): score = min(score, 18)
        scores.append(score)
    minimum = [850] * n if n != 2 else [2000, 3000]
    free = TOTAL_DXA - sum(minimum); total = sum(scores)
    result = [minimum[i] + int(free * scores[i] / total) for i in range(n)]
    result[-1] += TOTAL_DXA - sum(result)
    return result


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    repeat_header(table.rows[0])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = ""; shade(cell, C["light"])
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); inline(p, text, 9.5, C["navy"])
        for r in p.runs: r.bold = True
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(values):
            cell = cells[i]; cell.text = ""
            if row_index % 2: shade(cell, "FAFBFC")
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
            inline(p, text, 9.2)
    geometry(table, widths(headers, rows)); doc.add_paragraph().paragraph_format.space_after = Pt(1)


def callout(doc, text, risk=False):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"; no_split(table.rows[0])
    cell = table.cell(0, 0); shade(cell, C["red_bg"] if risk else C["gold_bg"])
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    inline(p, text, 10.5, C["red"] if risk else C["gold"]); geometry(table, [TOTAL_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def new_decimal_num(doc):
    root = doc.part.numbering_part.element
    abstracts = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    nums = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    aid = max(abstracts, default=-1) + 1; nid = max(nums, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(aid))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "decimal")
    txt = OxmlElement("w:lvlText"); txt.set(qn("w:val"), "%1.")
    jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr"); ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); ppr.append(ind)
    lvl.extend([start, fmt, txt, jc, ppr]); abstract.append(lvl); root.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(nid)); ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(aid)); num.append(ref); root.append(num)
    return nid


def new_bullet_num(doc):
    root = doc.part.numbering_part.element
    abstracts = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    nums = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    aid = max(abstracts, default=-1) + 1; nid = max(nums, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(aid))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "bullet")
    txt = OxmlElement("w:lvlText"); txt.set(qn("w:val"), "•")
    jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); ppr.append(ind)
    lvl.extend([start, fmt, txt, jc, ppr]); abstract.append(lvl); root.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(nid)); ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(aid)); num.append(ref); root.append(num)
    return nid


def numbered_paragraph(doc, text, num_id):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.2
    ppr = p._p.get_or_add_pPr(); numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id)); numpr.extend([ilvl, nid]); ppr.append(numpr)
    inline(p, text)


def setup(doc, preset, letter=False):
    section = doc.sections[0]
    if letter: section.page_width = Inches(8.5); section.page_height = Inches(11)
    else: section.page_width = Inches(8.267); section.page_height = Inches(11.693)
    section.left_margin = section.right_margin = Inches(.884 if not letter else 1)
    section.top_margin = section.bottom_margin = Inches(.75)
    section.header_distance = Inches(.35); section.footer_distance = Inches(.4); section.different_first_page_header_footer = True
    normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(C["ink"])
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25 if preset == "contract" else 1.10
    for style_name, size, color, before, after in (("Heading 1",16,C["blue"],14,8),("Heading 2",13,C["blue"],11,6),("Heading 3",12,C["dark"],8,4)):
        s = doc.styles[style_name]; s.font.name = "Calibri"; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
    bullet = doc.styles["List Bullet"]; bullet.paragraph_format.left_indent = Inches(.375); bullet.paragraph_format.first_line_indent = Inches(-.188); bullet.paragraph_format.space_after = Pt(4)


def page_number(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; font(p.add_run("Página "), 8.5, color=C["muted"])
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r); p._p.append(fld)


def furniture(doc, label, date):
    sec = doc.sections[0]; p = sec.header.paragraphs[0]; font(p.add_run(f"{label}  |  {date}"), 8.5, bold=True, color=C["muted"]); page_number(sec.footer.paragraphs[0])


def cover(doc, title, subtitle, kicker, metadata, warning):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(86); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(kicker.upper()), 10, bold=True, color=C["gold"])
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(title), 28, bold=True, color=C["navy"])
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(26); font(p.add_run(subtitle), 14, color=C["dark"])
    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2); table.style = "Table Grid"
        for i, (key, value) in enumerate(metadata):
            a,b = table.rows[i].cells; a.text = ""; b.text = ""; shade(a, C["gray"])
            font(a.paragraphs[0].add_run(key), 10, bold=True, color=C["muted"]); font(b.paragraphs[0].add_run(value), 10)
        geometry(table, [2200,7160])
    if warning:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(22); font(p.add_run(warning), 9.5, italic=True, color=C["red"])
    doc.add_page_break()


def parse(doc, text):
    lines = text.splitlines(); i = 0; skipped_h1 = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip(): i += 1; continue
        if line.startswith("# "):
            if not skipped_h1: skipped_h1 = True; i += 1; continue
            p = doc.add_paragraph(style="Heading 1"); inline(p, line[2:], 16, C["blue"]); i += 1; continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1"); inline(p, line[3:], 16, C["blue"]); i += 1; continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2"); inline(p, line[4:], 13, C["blue"]); i += 1; continue
        if line.startswith("> "):
            block=[]
            while i < len(lines) and lines[i].startswith("> "): block.append(lines[i][2:].strip()); i += 1
            callout(doc, " ".join(block)); continue
        if line.startswith("|") and i+1 < len(lines) and re.match(r"^\|?[\s:|-]+\|$", lines[i+1].strip()):
            headers=[x.strip() for x in line.strip().strip("|").split("|")]; i += 2; rows=[]
            while i < len(lines) and lines[i].strip().startswith("|"): rows.append([x.strip() for x in lines[i].strip().strip("|").split("|")]); i += 1
            add_table(doc, headers, rows); continue
        if re.match(r"^\s*\d+\.\s+", line):
            nid = new_decimal_num(doc)
            while i < len(lines):
                m = re.match(r"^\s*\d+\.\s+(.*)$", lines[i].rstrip())
                if not m: break
                numbered_paragraph(doc, m.group(1), nid); i += 1
            continue
        m = re.match(r"^\s*-\s+(.*)$", line)
        if m:
            nid = new_bullet_num(doc)
            while i < len(lines):
                m = re.match(r"^\s*-\s+(.*)$", lines[i].rstrip())
                if not m: break
                numbered_paragraph(doc, m.group(1), nid); i += 1
            continue
        parts=[line.strip()]; i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#",">","|")) and not re.match(r"^\s*(-|\d+\.)\s+", lines[i]): parts.append(lines[i].strip()); i += 1
        paragraph=" ".join(parts)
        if paragraph.startswith("**Não assinar") or paragraph.lower().startswith("**do not sign"): callout(doc, paragraph, True)
        elif not re.match(r"^\*\*(Data|Preço|Objetivo|Data-base|Localidade|Confiança)", paragraph):
            p=doc.add_paragraph(); inline(p,paragraph)


def main():
    ap=argparse.ArgumentParser(); ap.add_argument("input",type=Path); ap.add_argument("output",type=Path); ap.add_argument("--preset",choices=("contract","business"),default="business")
    ap.add_argument("--title"); ap.add_argument("--subtitle",default="Relatório profissional"); ap.add_argument("--kicker",default="Relatório"); ap.add_argument("--date",default=""); ap.add_argument("--metadata",action="append",default=[]); ap.add_argument("--warning",default=""); ap.add_argument("--letter",action="store_true")
    a=ap.parse_args(); source=a.input.read_text(encoding="utf-8"); inferred=next((x[2:] for x in source.splitlines() if x.startswith("# ")),a.input.stem); title=a.title or inferred
    metadata=[x.split("=",1) if "=" in x else (x,"") for x in a.metadata]
    doc=Document(); setup(doc,a.preset,a.letter); furniture(doc,title,a.date); cover(doc,title,a.subtitle,a.kicker,metadata,a.warning); parse(doc,source)
    doc.core_properties.title=title; doc.core_properties.author="Codex"; a.output.parent.mkdir(parents=True,exist_ok=True); doc.save(a.output); print(a.output)


if __name__ == "__main__": main()
